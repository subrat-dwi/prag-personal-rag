import logging
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

logger = logging.getLogger(__name__)

#-----------------docx parsing function------------------------
def parse_docx(path: str) -> str:
    """
    Extract text from a .docx file including paragraphs and tables.

    Args:
        path: Path to the .docx file.

    Returns:
        Extracted text as a single string.

    Raises:
        FileNotFoundError: If file doesn't exist.
        ValueError: If file is not a valid .docx.
        RuntimeError: If no text could be extracted.
    """
    docx_path = Path(path)

    if not docx_path.exists():
        raise FileNotFoundError(f"File not found: {path}")

    if docx_path.suffix.lower() != ".docx":
        raise ValueError(f"Not a .docx file: {path}")

    try:
        doc = Document(str(docx_path))
    except Exception as e:
        raise ValueError(f"Could not open .docx (possibly corrupt): {path}") from e

    blocks = []

    # extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            blocks.append(text)

    # extract tables — each cell on its own line, rows separated
    for table in doc.tables:
        for row in table.rows:
            row_cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_cells:
                blocks.append(" | ".join(row_cells))

    full_text = "\n".join(blocks).strip()

    if not full_text:
        raise RuntimeError(f"No text could be extracted from '{path}'.")

    logger.info("Extracted %d characters from '%s'", len(full_text), docx_path.name)
    return full_text